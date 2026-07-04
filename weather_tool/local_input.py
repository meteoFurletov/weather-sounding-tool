"""
Обработка зондирований из локальных файлов (Word .docx, обычный текст .txt или
сохранённая страница .html), скопированных вручную с нового сайта Вайоминга.

Зачем это нужно
---------------
С середины 2025 года архив радиозондирований Вайоминга переехал на новый сервер
(``weather.arcc.uwyo.edu``) и теперь отдаёт данные **только по одному сроку за
запрос**. Прежняя выгрузка списком за целый месяц (``fetcher.py``) больше не
работает, поэтому данные приходится копировать по срокам вручную.

Рекомендуемый порядок действий
------------------------------
1. На https://weather.arcc.uwyo.edu/upperair/sounding.shtml выберите дату, час,
   номер станции и тип вывода **"Text: List"**, затем нажмите Submit.
2. Выделите на странице весь результат — **обязательно вместе с заголовком**
   ``Observations for Station NNNNN at HH UTC DD Mon YYYY`` — и вставьте в
   документ Word. Повторите для каждого нужного срока в том же документе.
3. Передайте этот .docx (или .txt) в :func:`process_files`.

Дата и время каждого срока считываются из строки-заголовка автоматически.
Поддерживаются оба формата таблицы: новый (колонка ``SPED``, м/с; ``12 UTC``)
и старый (колонка ``SKNT``, узлы; ``12Z``).
"""

import os
import re

import numpy as np
import pandas as pd

from .logger import logger
from .processor import detect_inversions

# Метеорологические колонки таблицы зондирования — удаляются перед сохранением,
# в файл идут только вычисленные характеристики инверсии.
METEO_COLUMNS = [
    "PRES", "HGHT", "TEMP", "DWPT", "RELH", "MIXR",
    "DRCT", "SKNT", "SPED", "THTA", "THTE", "THTV",
]

_MONTHS = {
    m: i
    for i, m in enumerate(
        ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
         "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"],
        start=1,
    )
}

# Новый формат заголовка: "Observations for Station 26075 at 12 UTC 15 Jun 2025"
_NEW_HEADER = re.compile(
    r"Observations?\s+for\s+Station\s+(\d+)\s+at\s+(\d{1,2})\s*UTC\s+"
    r"(\d{1,2})\s+([A-Za-z]{3})[a-z]*\s+(\d{4})",
    re.IGNORECASE,
)

# Старый формат заголовка: "... Observations at 12Z 15 Jun 2025"
_OLD_HEADER = re.compile(
    r"(\d{2})Z\s+(\d{1,2})\s+([A-Za-z]{3})[a-z]*\s+(\d{4})",
    re.IGNORECASE,
)


def _to_datetime(hour, day, mon, year):
    """Собрать pandas.Timestamp без зависимости от локали (месяц по-английски)."""
    month = _MONTHS[mon[:3].title()]
    return pd.Timestamp(int(year), month, int(day), int(hour))


def read_document_text(path):
    """Прочитать текст из .docx, .txt или .html файла."""
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        try:
            import docx  # python-docx
        except ImportError as exc:
            raise ImportError(
                "Для чтения файлов Word нужен пакет python-docx. "
                "Установите его: pip install python-docx"
            ) from exc
        document = docx.Document(path)
        parts = [p.text for p in document.paragraphs]
        # На случай, если вставка превратилась в таблицу Word — читаем и её.
        for table in document.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    if ext == ".doc":
        raise ValueError(
            "Старый формат .doc не поддерживается. Пересохраните файл как .docx "
            "(Файл → Сохранить как → Документ Word .docx) или как обычный текст .txt."
        )

    if ext in (".htm", ".html"):
        from bs4 import BeautifulSoup

        with open(path, encoding="utf-8", errors="replace") as handle:
            soup = BeautifulSoup(handle.read(), "html.parser")
        return soup.get_text("\n")

    # .txt и всё остальное — читаем как обычный текст
    with open(path, encoding="utf-8", errors="replace") as handle:
        return handle.read()


def _find_headers(text):
    """Найти все заголовки сроков: список (позиция, время, номер станции)."""
    headers = []
    for match in _NEW_HEADER.finditer(text):
        station, hour, day, mon, year = match.groups()
        headers.append((match.start(), _to_datetime(hour, day, mon, year), int(station)))

    if not headers:
        # Ни одного нового заголовка — пробуем старый формат
        for match in _OLD_HEADER.finditer(text):
            hour, day, mon, year = match.groups()
            headers.append((match.start(), _to_datetime(hour, day, mon, year), None))

    headers.sort(key=lambda item: item[0])
    return headers


def _parse_table(block_text):
    """Разобрать текстовую таблицу зондирования в числовой DataFrame.

    Возвращает DataFrame с колонками PRES/HGHT/TEMP/... или ``None``, если
    таблица не найдена либо разметка по столбцам нарушена (например, при вставке
    в Word моноширинное выравнивание было потеряно).
    """
    lines = block_text.splitlines()

    # Ищем строку-заголовок таблицы (содержит PRES и HGHT)
    header_idx = None
    for idx, line in enumerate(lines):
        if "PRES" in line and "HGHT" in line:
            header_idx = idx
            break
    if header_idx is None:
        return None

    columns = lines[header_idx].split()

    rows = []
    for line in lines[header_idx + 1:]:
        stripped = line.strip()
        if not stripped:
            if rows:  # пустая строка после данных — конец таблицы
                break
            continue
        # Пропускаем разделители "----" и строку с единицами измерения (hPa, m, ...)
        if set(stripped) <= set("- ") or "hPa" in line:
            continue
        # Данные всегда начинаются с числа (давление). Если нет — таблица кончилась.
        if not re.match(r"-?\d", stripped):
            if rows:
                break
            continue
        # Фиксированная ширина колонки — 7 символов
        cells = [line[i : i + 7].strip() for i in range(0, len(line), 7)]
        rows.append(cells)

    if len(rows) < 2:
        return None

    # Приводим все строки к одинаковому числу колонок
    width = len(columns)
    norm = [(cells + [""] * width)[:width] for cells in rows]
    df = pd.DataFrame(norm, columns=columns)

    try:
        df = df.replace("", np.nan).astype(float)
    except ValueError:
        logger.warning(
            "Не удалось разобрать таблицу как числовую — возможно, при вставке в "
            "Word нарушено выравнивание столбцов. Скопируйте данные ещё раз как "
            "обычный текст (или сохраните срок в .txt)."
        )
        return None

    # Нормализуем колонку ветра: новый формат SPED (м/с) -> оставляем как есть,
    # старый SKNT (узлы) -> переводим в м/с. Для анализа инверсий ветер не нужен,
    # это лишь для единообразия вывода.
    if "SKNT" in df.columns and "SPED" not in df.columns:
        df["SKNT"] = (df["SKNT"] * 0.51444444444444).round(2)

    # Обязательные для анализа колонки
    if "TEMP" not in df.columns or "HGHT" not in df.columns:
        return None
    if df["TEMP"].notna().sum() < 2 or df["HGHT"].notna().sum() < 2:
        return None

    return df


def extract_soundings_from_text(text, default_station=None):
    """Извлечь все зондирования из текста.

    Возвращает список кортежей ``(DataFrame, время, номер_станции)``.
    """
    headers = _find_headers(text)
    soundings = []

    if not headers:
        logger.warning(
            "В файле не найдено ни одного заголовка срока "
            "('Observations for Station ... at HH UTC ...'). "
            "Убедитесь, что при копировании со страницы вы захватили строку "
            "с датой и временем над таблицей."
        )
        return soundings

    for position, obs_time, station in headers:
        end = len(text)
        for other_pos, _, _ in headers:
            if other_pos > position:
                end = min(end, other_pos)
        block = text[position:end]

        df = _parse_table(block)
        if df is None:
            logger.warning(f"Не удалось разобрать таблицу для срока {obs_time}.")
            continue
        soundings.append((df, obs_time, station if station is not None else default_station))

    logger.info(f"Извлечено сроков: {len(soundings)}")
    return soundings


def build_combined_file(all_inversions, combined_path="DATA.xlsx"):
    """Собрать сводный файл DATA.xlsx с несколькими листами анализа.

    Сетка дат строится по фактическому диапазону наблюдений (от самой ранней до
    самой поздней даты, шаг 12 ч), поэтому месяцы могут быть не подряд и данных
    может быть сколько угодно. Сроки, не попадающие в стандартную сетку 00/12,
    также сохраняются (объединение outer).
    """
    combined = pd.concat(all_inversions, ignore_index=True)
    combined["date"] = pd.to_datetime(combined["date"])

    start = combined["date"].min().normalize()
    end = combined["date"].max().normalize() + pd.Timedelta(days=1)
    grid = pd.date_range(start=start, end=end, freq="12h")
    grid_df = pd.DataFrame({"date": grid})

    data_full = (
        grid_df.merge(combined, on="date", how="outer")
        .sort_values("date")
        .reset_index(drop=True)
    )

    subsets = {
        "df_full": data_full,
        "df_ground": data_full[data_full["Ground"] == 1],
        "df_not_ground": data_full[data_full["Ground"] == 0],
        "df_day": data_full[data_full["Day"] == 1],
        "df_night": data_full[data_full["Night"] == 1],
        "df_ground_night": data_full[
            (data_full["Ground"] == 1) & (data_full["Night"] == 1)
        ],
        "df_ground_day": data_full[
            (data_full["Ground"] == 1) & (data_full["Day"] == 1)
        ],
        "df_not_ground_night": data_full[
            (data_full["Ground"] == 0) & (data_full["Night"] == 1)
        ],
        "df_not_ground_day": data_full[
            (data_full["Ground"] == 0) & (data_full["Day"] == 1)
        ],
    }

    with pd.ExcelWriter(combined_path, engine="xlsxwriter") as writer:
        for name, frame in subsets.items():
            frame.to_excel(writer, sheet_name=name, index=False)

    logger.info(
        f"Создан сводный файл {combined_path} ({len(subsets)} листов анализа)"
    )


def process_files(paths, station=None, output_dir=None, combined_path="DATA.xlsx"):
    """Обработать один или несколько локальных файлов (.docx/.txt/.html).

    Параметры
    ---------
    paths : str | list[str]
        Путь к файлу или список путей.
    station : int, optional
        Номер станции. Используется в имени папки с результатами и как запасной
        вариант, если в файле его не удалось определить.
    output_dir : str, optional
        Папка для отдельных файлов по срокам. По умолчанию
        ``soundings_<год>_<станция>``.
    combined_path : str
        Имя сводного файла. По умолчанию ``DATA.xlsx``.
    """
    if isinstance(paths, str):
        paths = [paths]

    texts = []
    for path in paths:
        if not os.path.exists(path):
            logger.error(f"Файл не найден: {path}")
            continue
        logger.info(f"Читаю файл: {path}")
        texts.append(read_document_text(path))

    return process_text(
        "\n".join(texts),
        station=station,
        output_dir=output_dir,
        combined_path=combined_path,
    )


def process_text(text, station=None, output_dir=None, combined_path="DATA.xlsx"):
    """Обработать сырой текст со сроками (из файла или скачанный с сайта).

    Ищет в тексте зондирования (по строке-заголовку), считает инверсии, сохраняет
    отдельные файлы по срокам и сводный файл. Возвращает словарь-итог.
    """
    soundings = extract_soundings_from_text(text, default_station=station)

    summary = {"processed": 0, "failed": 0, "files": []}
    all_inversions = []

    if not soundings:
        logger.error("Не найдено ни одного зондирования для обработки.")
        return summary

    # Определяем станцию и год для имени папки результатов
    dates = [obs_time for _, obs_time, _ in soundings]
    stations = [st for _, _, st in soundings if st is not None]
    year = min(dates).year
    resolved_station = station or (stations[0] if stations else "station")
    if output_dir is None:
        output_dir = f"soundings_{year}_{resolved_station}"
    os.makedirs(output_dir, exist_ok=True)

    for df, obs_time, _ in soundings:
        df_inverse = detect_inversions(df, obs_time)

        if df_inverse is not None and not df_inverse.empty:
            summary["processed"] += 1

            single_data = df_inverse.drop(METEO_COLUMNS, axis=1, errors="ignore")
            single_data = single_data[single_data["ΔT"] > 0].drop_duplicates()

            timestamp = obs_time.strftime("%Y%m%d_%H%M")
            filename = f"{output_dir}/DATA_{timestamp}.xlsx"
            with pd.ExcelWriter(filename, engine="xlsxwriter") as writer:
                single_data.to_excel(writer, sheet_name="inversion_data")

            logger.info(f"Сохранено: {filename}")
            summary["files"].append(filename)
            all_inversions.append(single_data)
        else:
            summary["failed"] += 1

    if all_inversions:
        build_combined_file(all_inversions, combined_path)

    return summary
